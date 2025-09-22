"""DOC parser with DOCX conversion for full-featured parsing."""

from pathlib import Path
from typing import Set, List, Optional
import tempfile
import os

try:
    import docx2txt
    from docx import Document as DocxDocument
    DOC_SUPPORT = True
except ImportError:
    DOC_SUPPORT = False

from ...core.interfaces import BaseParser
from ...core.document import Document
from ...core.logging import get_logger
from ...data.structures.metadata import Metadata
from ...data.structures.document.word import DocDocument as DocumentorDocDocument
from ...data.structures.fragment.text import ParagraphFragment
from .docx_parser import DocxParser
from .word_com_converter import WordCOMConverter

logger = get_logger(__name__)


class DocParser(BaseParser):
    """
    Parser for DOC files using Microsoft Word COM conversion to DOCX.
    
    This parser converts DOC files to DOCX format using Microsoft Word COM automation
    (Windows only), then uses the full-featured DOCX parser to extract rich content
    including formatting, images, tables, and structure.
    """
    
    def __init__(self):
        """Initialize DOC parser."""
        if not DOC_SUPPORT:
            raise ImportError("docx2txt and python-docx libraries are required for DOC processing")
        
        # Initialize DOCX parser for conversion-based processing
        self.docx_parser = DocxParser()
        
        # Initialize Word COM converter with test
        try:
            self.word_com_converter = WordCOMConverter()
            if self.word_com_converter.is_available:
                # Test actual COM access
                if self._test_word_com_access():
                    self.word_com_available = True
                    logger.info("DOC parser initialized with Word COM conversion support")
                else:
                    self.word_com_available = False
                    self.word_com_converter = None
                    logger.warning("Word COM test failed, using fallback only")
            else:
                self.word_com_available = False
                self.word_com_converter = None
                logger.warning("Word COM not available, using fallback only")
        except Exception as e:
            self.word_com_available = False
            self.word_com_converter = None
            logger.warning(f"Word COM initialization failed: {e}")
    
    def supported_extensions(self) -> Set[str]:
        """Get supported extensions."""
        return {'.doc'}
    
    def parse(self, file_path: Path) -> Document:
        """
        Parse DOC file into a Document object using DOCX conversion.
        
        Args:
            file_path: Path to DOC file
            
        Returns:
            Document: Parsed document with rich fragments
            
        Raises:
            FileNotFoundError: If file not found
            ValueError: If file cannot be processed
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not file_path.is_file():
            raise ValueError(f"Path is not a file: {file_path}")
        
        if file_path.suffix.lower() not in self.supported_extensions():
            raise ValueError(f"Unsupported file extension: {file_path.suffix}")
        
        logger.info(f"Parsing DOC file with conversion: {file_path}")
        
        try:
            # Try Word COM conversion first
            if self.word_com_available:
                docx_path = self._convert_with_word_com(file_path)
            else:
                # Fallback to basic conversion
                docx_path = self._convert_doc_to_docx(file_path)
            
            try:
                # Parse the converted DOCX file using the full DOCX parser
                document = self.docx_parser.parse(docx_path)
                
                # Update metadata to reflect original DOC source
                if hasattr(document, 'metadata'):
                    if self.word_com_available:
                        document.metadata.processing_method = "doc_processing_with_word_com"
                    else:
                        document.metadata.processing_method = "doc_processing_fallback"
                    document.metadata.original_format = "doc"
                    document.metadata.converted_from = str(file_path)
                
                logger.info(f"Successfully parsed DOC file via conversion: {len(document.fragments())} fragments")
                return document
                
            finally:
                # Clean up temporary DOCX file
                self._cleanup_temp_file(docx_path)
                
        except Exception as e:
            logger.error(f"Error processing DOC file {file_path}: {e}")
            # Fallback to basic text extraction
            return self._fallback_text_extraction(file_path)
    
    def _convert_with_word_com(self, doc_path: Path) -> Path:
        """
        Convert DOC file to DOCX using Word COM.
        
        Args:
            doc_path: Path to input DOC file
            
        Returns:
            Path to temporary DOCX file
        """
        if not self.word_com_available:
            raise RuntimeError("Word COM converter not available")
        
        # Create temporary directory for conversion
        temp_dir = Path(tempfile.gettempdir()) / f"documentor_word_com_{os.getpid()}"
        temp_dir.mkdir(exist_ok=True)
        
        try:
            # Convert using Word COM
            docx_path = self.word_com_converter.convert_doc_to_docx(doc_path, temp_dir)
            return docx_path
        except Exception as e:
            # Clean up temp directory on error
            self.word_com_converter.cleanup_temp_files(temp_dir)
            raise e
    
    def _convert_doc_to_docx(self, doc_path: Path) -> Path:
        """
        Fallback conversion when no proper converters are available.
        
        Creates a basic DOCX with error message.
        
        Args:
            doc_path: Path to input DOC file
            
        Returns:
            Path to temporary DOCX file
        """
        # Create new DOCX document with error message
        docx_doc = DocxDocument()
        docx_doc.add_paragraph(f"DOC file: {doc_path.name}")
        docx_doc.add_paragraph("Error: No conversion tools available (LibreOffice or Word COM)")
        docx_doc.add_paragraph("Please install LibreOffice or Microsoft Word for full DOC support.")
        
        # Create temporary file
        temp_dir = Path(tempfile.gettempdir()) / "documentor_doc_fallback"
        temp_dir.mkdir(exist_ok=True)
        temp_docx_path = temp_dir / f"{doc_path.stem}_fallback_{os.getpid()}.docx"
        
        # Save DOCX file
        docx_doc.save(str(temp_docx_path))
        
        logger.warning(f"Created fallback DOCX for DOC file: {temp_docx_path}")
        return temp_docx_path
    
    def _cleanup_temp_file(self, file_path: Path) -> None:
        """
        Clean up temporary file.
        
        Args:
            file_path: Path to temporary file to delete
        """
        try:
            if file_path.exists():
                file_path.unlink()
                logger.debug(f"Cleaned up temporary file: {file_path}")
        except Exception as e:
            logger.warning(f"Failed to clean up temporary file {file_path}: {e}")
    
    def _fallback_text_extraction(self, file_path: Path) -> Document:
        """
        Fallback to basic text extraction if all conversion methods fail.
        
        Args:
            file_path: Path to DOC file
            
        Returns:
            Document with error message fragments
        """
        logger.warning(f"All conversion methods failed for: {file_path}")
        
        metadata = Metadata.from_path(
            file_path,
            processing_method="doc_processing_fallback"
        )
        
        fragments = [
            ParagraphFragment(value=f"DOC file: {file_path.name}"),
            ParagraphFragment(value="Error: All conversion methods failed"),
            ParagraphFragment(value="Please install LibreOffice or Microsoft Word for DOC support")
        ]
        
        document = DocumentorDocDocument(
            fragments=fragments,
            metadata=metadata
        )
        
        logger.info(f"Fallback extraction created {len(fragments)} fragments from DOC file")
        return document
    
    def _test_word_com_access(self) -> bool:
        """
        Test actual Word COM access by creating and closing a document.
        
        Returns:
            bool: True if COM access works, False otherwise
        """
        try:
            import win32com.client
            word = None
            try:
                # Test COM access
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                word.DisplayAlerts = False
                
                # Test creating a document
                doc = word.Documents.Add()
                doc.Close(SaveChanges=False)
                
                logger.debug("Word COM access test successful")
                return True
                
            finally:
                if word:
                    word.Quit()
                    
        except Exception as e:
            logger.debug(f"Word COM access test failed: {e}")
            return False