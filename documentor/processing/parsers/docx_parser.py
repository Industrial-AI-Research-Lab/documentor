"""DOCX parser for Microsoft Word documents."""

from pathlib import Path
from typing import Set, List, Dict, Any, Optional, Iterator
import logging

try:
    from docx import Document as PythonDocxDocument
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

from ...core.interfaces import BaseParser
from ...core.document import Document
from ...core.logging import get_logger
from ...data.structures.metadata import Metadata
from ...data.structures.document.word import DocxDocument as DocumentorDocxDocument
from ...data.structures.fragment import (
    ParagraphFragment, 
    HeaderFragment,
    TitleFragment,
    TableFragment,
    ListItemFragment,
    FootnoteFragment,
    ImageFragment,
    HyperlinkFragment,
    StyleFragment,
    CommentFragment,
    BreakFragment
)

logger = get_logger(__name__)


class DocxParser(BaseParser):
    """
    Parser for DOCX files using python-docx library.
    
    Extracts rich structure including:
    - Paragraphs with formatting
    - Headers and titles 
    - Tables
    - Lists (numbered and bulleted)
    - Images
    - Hyperlinks
    - Footnotes
    - Comments
    - Page/section breaks
    """
    
    def __init__(self):
        """Initialize DOCX parser."""
        try:
            import docx
            logger.info("DOCX parser initialized")
        except ImportError as e:
            raise ImportError("python-docx library is required for DOCX processing") from e
    
    def supported_extensions(self) -> Set[str]:
        """Get supported extensions."""
        return {'.docx'}
    
    def parse(self, file_path: Path) -> Document:
        """
        Parse DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            DocumentorDocxDocument: Document with fragments
            
        Raises:
            FileNotFoundError: If file not found
            ValueError: If file cannot be processed
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not file_path.is_file():
            raise ValueError(f"Path is not a file: {file_path}")
        
        if file_path.suffix.lower() not in self.supported_extensions():
            raise ValueError(f"Unsupported file extension: {file_path.suffix}")
        
        logger.info(f"Parsing DOCX file: {file_path}")
        
        try:
            # Load DOCX document
            docx_doc = PythonDocxDocument(str(file_path))
            
            # Extract all fragments
            all_fragments = []
            
            # Extract document properties
            doc_props = self._extract_document_properties(docx_doc)
            
            # Extract main document content (with images inline)
            fragments = list(self._extract_document_content(docx_doc, file_path))
            all_fragments.extend(fragments)
            
            # Extract Table of Contents if present
            toc_fragments = list(self._extract_toc_fields(docx_doc))
            all_fragments.extend(toc_fragments)
            
            # Extract headers and footers
            header_footer_fragments = list(self._extract_headers_footers(docx_doc))
            all_fragments.extend(header_footer_fragments)
            
            # Extract footnotes
            footnote_fragments = list(self._extract_footnotes(docx_doc))
            all_fragments.extend(footnote_fragments)
            
            # Extract comments
            comment_fragments = list(self._extract_comments(docx_doc))
            all_fragments.extend(comment_fragments)
            
            # Create document structure and relationships
            self._detect_and_process_toc(all_fragments)
            self._create_document_hierarchy(all_fragments)
            self._link_images_with_captions(all_fragments)
            self._link_citations_with_references(all_fragments)
            
            # Create metadata
            metadata = Metadata.from_path(
                file_path,
                processing_method="docx_processing"
            )
            
            # Add DOCX-specific metadata
            metadata.params.update(doc_props)
            
            # Count stats
            word_count = sum(len(f.value.split()) for f in all_fragments 
                           if hasattr(f, 'value') and isinstance(f.value, str))
            char_count = sum(len(f.value) for f in all_fragments 
                           if hasattr(f, 'value') and isinstance(f.value, str))
            
            # Count image fragments
            image_count = sum(1 for f in all_fragments if f.__class__.__name__ == 'ImageFragment')
            
            logger.info(f"Extracted {len(all_fragments)} fragments, {word_count} words, {image_count} images in document order")
            
            # Create DocumentorDocxDocument
            document = DocumentorDocxDocument(
                fragments=all_fragments,
                metadata=metadata,
                pages_count=doc_props.get('pages_count', 1),
                word_count=word_count,
                character_count=char_count,
                document_properties=doc_props,
                sections_count=len(docx_doc.sections),
                has_headers_footers=len(header_footer_fragments) > 0,
                has_footnotes=len(footnote_fragments) > 0,
                has_comments=len(comment_fragments) > 0
            )
            
            # Build document structure from headers
            # document.create_structure_from_headers()  # This method doesn't exist yet
            
            return document
            
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {e}")
            raise ValueError(f"Failed to process DOCX file: {e}")
    
    def _extract_document_properties(self, docx_doc: "PythonDocxDocument") -> Dict[str, Any]:
        """Extract document properties and metadata."""
        props = {}
        
        try:
            core_props = docx_doc.core_properties
            props.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'last_modified_by': core_props.last_modified_by or '',
                'created': core_props.created.isoformat() if core_props.created else None,
                'modified': core_props.modified.isoformat() if core_props.modified else None,
                'revision': core_props.revision or 0
            })
        except Exception as e:
            logger.warning(f"Error extracting core properties: {e}")
        
        # Estimate page count (rough approximation)
        try:
            total_paragraphs = len(docx_doc.paragraphs)
            props['estimated_pages'] = max(1, total_paragraphs // 25)  # ~25 paragraphs per page
        except Exception:
            props['estimated_pages'] = 1
        
        return props
    
    def _extract_document_content(self, docx_doc: "PythonDocxDocument", file_path) -> Iterator:
        """
        Extract document content in strict sequential order.
        
        Processes paragraphs, tables, and images while preserving their original sequence.
        Extracts comprehensive metadata including formatting, structure, and hierarchy.
        Handles multiple image formats: DrawingML, VML, Objects, and embedded elements.
        
        Args:
            docx_doc: Python-docx Document object
            file_path: Path to DOCX file for direct ZIP archive access
            
        Yields:
            Fragment objects with rich metadata in document order
        """
        paragraph_count = 0
        current_page = 1
        estimated_pages = max(1, len(docx_doc.paragraphs) // 15)
        image_counter = 0
        processed_image_hashes = set()
        
        # Create relationships mapping for all document elements
        relationships = {}
        try:
            for rel_id, rel in docx_doc.part.rels.items():
                relationships[rel_id] = rel
                
                # Log only image relationships for clarity
                if hasattr(rel, 'target_ref') and 'media/' in rel.target_ref:
                    logger.info(f"Found image: {rel_id} -> {rel.target_ref}")
                    
            image_count = len([r for r in relationships.values() if hasattr(r, 'target_ref') and 'media/' in r.target_ref])
            logger.info(f"Loaded {image_count} image relationships")
        except Exception as e:
            logger.warning(f"Error loading relationships: {e}")
        
        # Process each paragraph in document order
        for para in docx_doc.paragraphs:
            # Check for page breaks in paragraph
            if self._has_page_break(para._element):
                current_page += 1
                logger.debug(f"Found page break, now on page {current_page}")
            
            # Better page estimation
            if current_page == 1 and paragraph_count > 0:
                if paragraph_count > 10:
                    estimated_page = min(2 + (paragraph_count - 10) // 15, estimated_pages)
                else:
                    estimated_page = 1
            else:
                estimated_page = current_page
            
            # Process paragraph in document order - text AND images in sequence
            for result in self._process_paragraph_in_document_order(
                para, paragraph_count, estimated_page, image_counter, processed_image_hashes, relationships, file_path
            ):
                if isinstance(result, tuple) and result[0] == 'image':
                    image_counter = result[2]  # Update counter
                    yield result[1]  # Yield image fragment
                else:
                    yield result  # Yield text fragment
            
            paragraph_count += 1
        
        # Process tables separately (they may contain images)
        for table in docx_doc.tables:
            estimated_page = current_page if current_page > 1 else max(1, paragraph_count // 15)
            yield from self._process_table_with_images(table, estimated_page, image_counter, processed_image_hashes, relationships, file_path)
        
        # All images should now be extracted in document order during paragraph processing
        logger.info(f"Document processing complete. Total image counter: {image_counter}")
    
    def _has_page_break(self, element) -> bool:
        """Check if element contains a page break."""
        try:
            # Look for page breaks in the element
            page_breaks = element.xpath('.//w:br[@w:type="page"]')
            if len(page_breaks) > 0:
                return True
                
            # Also check for section breaks that often indicate new pages
            section_breaks = element.xpath('.//w:sectPr')
            if len(section_breaks) > 0:
                return True
                
            return False
        except Exception:
            return False
    
    def _estimate_page_from_position(self, paragraph_index: int, total_paragraphs: int, estimated_pages: int) -> int:
        """Estimate page number based on paragraph position."""
        if total_paragraphs == 0:
            return 1
        
        # Simple estimation: distribute paragraphs evenly across pages
        paragraphs_per_page = max(1, total_paragraphs / estimated_pages)
        estimated_page = max(1, int(paragraph_index / paragraphs_per_page) + 1)
        
        return min(estimated_page, estimated_pages)
    
    def _process_paragraph(self, paragraph, para_index: int, page: int = 1) -> Iterator:
        """Process a single paragraph."""
        text = paragraph.text.strip()
        if not text:
            return
        
        # Determine paragraph type based on style
        style_name = paragraph.style.name if paragraph.style else ""
        
        # Check if it's a header
        if self._is_header_style(style_name):
            level = self._get_header_level(style_name)
            if level == 1:
                fragment = TitleFragment(
                    value=text,
                    id=f"title_{para_index}",
                    level=level,
                    page=page
                )
            else:
                fragment = HeaderFragment(
                    value=text,
                    id=f"header_{para_index}_{level}",
                    level=level,
                    page=page
                )
        
        # Check if it's a list item
        elif self._is_list_item(paragraph):
            list_info = self._get_list_info(paragraph)
            fragment = ListItemFragment(
                value=text,
                id=f"list_item_{para_index}",
                list_type=list_info['type'],
                list_level=list_info['level'],
                marker=list_info['marker'],
                page=page
            )
        
        # Regular paragraph
        else:
            # Check for rich formatting
            if self._has_rich_formatting(paragraph):
                formatting = self._extract_formatting(paragraph)
                # Filter formatting to only include StyleFragment fields
                style_formatting = {
                    'font_family': formatting.get('font_family'),
                    'font_size': formatting.get('font_size'),
                    'is_bold': formatting.get('is_bold', False),
                    'is_italic': formatting.get('is_italic', False),
                    'is_underline': formatting.get('is_underline', False),
                    'text_color': formatting.get('text_color'),
                    'background_color': formatting.get('background_color'),
                    'alignment': formatting.get('alignment')
                }
                # Remove None values
                style_formatting = {k: v for k, v in style_formatting.items() if v is not None}
                
                fragment = StyleFragment(
                    value=text,
                    id=f"styled_para_{para_index}",
                    page=page,
                    **style_formatting
                )
            else:
                fragment = ParagraphFragment(
                    value=text,
                    id=f"para_{para_index}",
                    page=page
                )
        
        # Add detailed formatting metadata to fragment
        if fragment:
            detailed_formatting = self._extract_formatting(paragraph)
            if not hasattr(fragment, 'metadata'):
                fragment.metadata = type('obj', (object,), {})()
            
            # Add all formatting metadata
            for key, value in detailed_formatting.items():
                if value is not None:
                    setattr(fragment.metadata, key, value)
        
        # Extract hyperlinks in paragraph
        hyperlinks = list(self._extract_hyperlinks(paragraph, para_index))
        if hyperlinks:
            yield from hyperlinks
        else:
            yield fragment
        
        # Extract images in paragraph
        yield from self._extract_paragraph_images(paragraph, para_index, page)
    
    def _process_table(self, table, page: int = 1) -> Iterator:
        """Process a table."""
        # Simple table processing - can be enhanced
        table_data = []
        
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data.append(row_data)
        
        if table_data:
            fragment = TableFragment(
                value=table_data,
                id=f"table_{id(table)}",
                page=page
            )
            
            # Add table metadata
            if not hasattr(fragment, 'metadata'):
                fragment.metadata = type('obj', (object,), {})()
            
            fragment.metadata.row_count = len(table.rows)
            fragment.metadata.column_count = len(table.rows[0].cells) if table.rows else 0
            fragment.metadata.total_cells = fragment.metadata.row_count * fragment.metadata.column_count
            fragment.metadata.has_header = True  # Assume first row is header
            fragment.metadata.table_style = getattr(table.style, 'name', None) if table.style else None
            
            # Calculate text statistics
            total_text = ' '.join(' '.join(row) for row in table_data)
            fragment.metadata.text_length = len(total_text)
            fragment.metadata.word_count = len(total_text.split())
            
            yield fragment
    
    def _extract_hyperlinks(self, paragraph, para_index: int) -> Iterator:
        """Extract hyperlinks from paragraph."""
        hyperlink_count = 0
        
        # Look for hyperlinks in paragraph XML
        for hyperlink in paragraph._element.xpath('.//w:hyperlink'):
            link_text = ""
            link_url = ""
            
            # Extract text
            for run in hyperlink.xpath('.//w:t'):
                if run.text:
                    link_text += run.text
            
            # Extract URL from relationship
            r_id = hyperlink.get(qn('r:id'))
            if r_id:
                try:
                    link_url = paragraph.part.rels[r_id].target_ref
                except Exception:
                    pass
            
            if link_text.strip():
                fragment = HyperlinkFragment(
                    value=link_text.strip(),
                    url=link_url,
                    id=f"hyperlink_{para_index}_{hyperlink_count}"
                )
                yield fragment
                hyperlink_count += 1
    
    def _extract_headers_footers(self, docx_doc: "PythonDocxDocument") -> Iterator:
        """Extract headers and footers."""
        # This is a simplified implementation
        # Full implementation would need to access headers/footers XML
        return iter([])
    
    def _extract_footnotes(self, docx_doc: "PythonDocxDocument") -> Iterator:
        """Extract footnotes."""
        # This is a simplified implementation
        # Full implementation would need to access footnotes XML
        return iter([])
    
    def _extract_comments(self, docx_doc: "PythonDocxDocument") -> Iterator:
        """Extract comments."""
        # This is a simplified implementation  
        # Full implementation would need to access comments XML
        return iter([])
    
    def _extract_toc_fields(self, docx_doc: "PythonDocxDocument") -> Iterator:
        """Extract Table of Contents fields from document."""
        try:
            logger.info("Searching for TOC fields in document")
            toc_found = False
            
            # Try multiple approaches to find TOC
            
            # Approach 1: Look for TOC field codes
            for para_idx, para in enumerate(docx_doc.paragraphs):
                para_xml = para._element
                
                # Look for TOC field codes
                field_codes = para_xml.xpath('.//w:instrText')
                for field_code in field_codes:
                    if field_code.text and 'TOC' in field_code.text:
                        toc_found = True
                        logger.info(f"Found TOC field code in paragraph {para_idx}: {field_code.text}")
                        
                        text = para.text.strip()
                        if text:
                            fragment = ParagraphFragment(
                                value=text,
                                id=f"toc_field_{para_idx}",
                                page=2
                            )
                            
                            if not hasattr(fragment, 'metadata'):
                                fragment.metadata = type('obj', (object,), {})()
                            fragment.metadata.is_toc_field = True
                            fragment.metadata.field_type = 'TOC'
                            fragment.metadata.toc_field_code = field_code.text
                            
                            logger.info(f"Created TOC field fragment: {text}")
                            yield fragment
                
                # Look for result text of TOC fields (fldChar end)
                fld_chars = para_xml.xpath('.//w:fldChar[@w:fldCharType="end"]')
                if fld_chars:
                    # Check if the previous paragraph had a TOC begin
                    if para_idx > 0:
                        prev_para = docx_doc.paragraphs[para_idx - 1]._element
                        toc_begins = prev_para.xpath('.//w:fldChar[@w:fldCharType="begin"]')
                        if toc_begins:
                            text = para.text.strip()
                            if text:
                                toc_found = True
                                logger.info(f"Found TOC result text at paragraph {para_idx}: {text}")
                                
                                fragment = ParagraphFragment(
                                    value=text,
                                    id=f"toc_result_{para_idx}",
                                    page=2
                                )
                                
                                if not hasattr(fragment, 'metadata'):
                                    fragment.metadata = type('obj', (object,), {})()
                                fragment.metadata.is_toc_result = True
                                
                                yield fragment
                
                # Look for special TOC styles
                if para.style and para.style.name:
                    style_name = para.style.name.lower()
                    if 'toc' in style_name or 'contents' in style_name:
                        text = para.text.strip()
                        if text:
                            toc_found = True
                            logger.info(f"Found TOC by style in paragraph {para_idx}: {style_name} -> {text}")
                            
                            fragment = ParagraphFragment(
                                value=text,
                                id=f"toc_style_{para_idx}",
                                page=2
                            )
                            
                            if not hasattr(fragment, 'metadata'):
                                fragment.metadata = type('obj', (object,), {})()
                            fragment.metadata.is_toc_entry = True
                            fragment.metadata.toc_style = para.style.name
                            
                            yield fragment
                
                # Look for hyperlink fields that might be TOC entries
                hyperlinks = para_xml.xpath('.//w:hyperlink')
                for hyperlink in hyperlinks:
                    # Check if this is an internal link (TOC entry)
                    anchor = hyperlink.get(qn('w:anchor'))
                    if anchor:
                        text = ""
                        for text_run in hyperlink.xpath('.//w:t'):
                            text += text_run.text
                        
                        if text.strip():
                            toc_found = True
                            logger.info(f"Found TOC hyperlink in paragraph {para_idx}: {text.strip()} -> {anchor}")
                            
                            # This looks like a TOC entry
                            fragment = ParagraphFragment(
                                value=text.strip(),
                                id=f"toc_hyperlink_{para_idx}",
                                page=2  # TOC is typically on page 2
                            )
                            
                            # Add TOC metadata
                            if not hasattr(fragment, 'metadata'):
                                fragment.metadata = type('obj', (object,), {})()
                            fragment.metadata.is_toc_entry = True
                            fragment.metadata.toc_anchor = anchor
                            fragment.metadata.is_internal_link = True
                            
                            yield fragment
            
            if not toc_found:
                logger.warning("No TOC fields found in document")
                            
        except Exception as e:
            logger.warning(f"Error extracting TOC fields: {e}")
            return iter([])
    
    def _create_document_hierarchy(self, fragments: List) -> None:
        """Create hierarchical structure based on headers and titles."""
        current_section = None
        current_subsection = None
        
        for i, fragment in enumerate(fragments):
            # Skip non-text fragments
            if not hasattr(fragment, 'level'):
                continue
                
            fragment_type = fragment.__class__.__name__
            
            if fragment_type == 'TitleFragment':
                # Start new section
                current_section = {
                    'title': fragment.value,
                    'level': fragment.level,
                    'start_index': i,
                    'subsections': [],
                    'content_fragments': []
                }
                if not hasattr(fragment, 'metadata'):
                    fragment.metadata = type('obj', (object,), {})()
                fragment.metadata.section = current_section['title']
                fragment.metadata.is_section_start = True
                current_subsection = None
                
            elif fragment_type == 'HeaderFragment':
                if current_section:
                    # Add subsection
                    current_subsection = {
                        'title': fragment.value,
                        'level': fragment.level,
                        'start_index': i,
                        'content_fragments': []
                    }
                    current_section['subsections'].append(current_subsection)
                    if not hasattr(fragment, 'metadata'):
                        fragment.metadata = type('obj', (object,), {})()
                    fragment.metadata.section = current_section['title']
                    fragment.metadata.subsection = current_subsection['title']
                    fragment.metadata.is_subsection_start = True
            
            else:
                # Regular content fragment
                if current_section:
                    if not hasattr(fragment, 'metadata'):
                        fragment.metadata = type('obj', (object,), {})()
                    fragment.metadata.section = current_section['title']
                    
                    if current_subsection:
                        fragment.metadata.subsection = current_subsection['title']
                        current_subsection['content_fragments'].append(i)
                    else:
                        current_section['content_fragments'].append(i)
    
    def _link_images_with_captions(self, fragments: List) -> None:
        """Link images with their captions (Рис. X. описание)."""
        import re
        
        # Pattern for figure captions: "Рис. N. описание" or "Рисунок N. описание"
        caption_pattern = re.compile(r'^(?:Рис|Рисунок)\.?\s*(\d+)\.?\s*(.+)', re.IGNORECASE)
        
        logger.info(f"Looking for image-caption links in {len(fragments)} fragments")
        
        # First pass: find all captions
        captions = {}
        for i, fragment in enumerate(fragments):
            if hasattr(fragment, 'value') and isinstance(fragment.value, str):
                match = caption_pattern.match(fragment.value.strip())
                if match:
                    figure_number = match.group(1)
                    caption_text = match.group(2)
                    captions[figure_number] = {
                        'text': caption_text,
                        'fragment_index': i,
                        'fragment': fragment
                    }
                    logger.info(f"Found caption: Рис. {figure_number} at fragment {i}")
        
        # Second pass: link images with captions in document order
        image_fragments = []
        for i, fragment in enumerate(fragments):
            if fragment.__class__.__name__ == 'ImageFragment':
                image_fragments.append((i, fragment))
                logger.info(f"Found image fragment at index {i}")
        
        # Link images with captions based on proximity in document
        for img_idx, (i, fragment) in enumerate(image_fragments):
            logger.info(f"Looking for caption near image at position {i}")
            
            # Find the nearest caption after this image (within next 5 fragments)
            best_caption = None
            best_distance = float('inf')
            
            for fig_num, caption_info in captions.items():
                j = caption_info['fragment_index']
                
                # Only consider captions that come after the image (j > i) and within reasonable distance
                if j > i and (j - i) <= 5:
                    distance = j - i
                    if distance < best_distance:
                        best_distance = distance
                        best_caption = (fig_num, caption_info)
                        logger.info(f"Found candidate caption at position {j} (distance {distance}): Рис. {fig_num}")
            
            if best_caption:
                figure_number, caption_info = best_caption
                caption_fragment = caption_info['fragment']
                caption_text = caption_info['text']
                j = caption_info['fragment_index']
                
                # Add metadata to image
                if not hasattr(fragment, 'metadata'):
                    fragment.metadata = type('obj', (object,), {})()
                fragment.metadata.figure_number = figure_number
                fragment.metadata.caption = caption_text
                fragment.metadata.caption_fragment_id = getattr(caption_fragment, 'id', f'fragment_{j:03d}')
                
                # Add metadata to caption
                if not hasattr(caption_fragment, 'metadata'):
                    caption_fragment.metadata = type('obj', (object,), {})()
                caption_fragment.metadata.is_figure_caption = True
                caption_fragment.metadata.figure_number = figure_number
                caption_fragment.metadata.linked_image_fragment_id = getattr(fragment, 'id', f'fragment_{i:03d}')
                
                # Create hierarchical relationship: Image -> Caption (child)
                # Set image as parent of caption
                caption_fragment.metadata.parent_fragment_id = getattr(fragment, 'id', f'fragment_{i:03d}')
                caption_fragment.metadata.hierarchy_type = 'image_caption'
                caption_fragment.metadata.hierarchy_level = 1  # Caption is level 1 under image
                
                # Set caption as child of image
                if not hasattr(fragment, 'metadata'):
                    fragment.metadata = type('obj', (object,), {})()
                if not hasattr(fragment.metadata, 'child_fragment_ids'):
                    fragment.metadata.child_fragment_ids = []
                fragment.metadata.child_fragment_ids.append(getattr(caption_fragment, 'id', f'fragment_{j:03d}'))
                fragment.metadata.hierarchy_type = 'image_with_caption'
                fragment.metadata.hierarchy_level = 0  # Image is root level
                
                logger.info(f"Created hierarchy: Image {i} -> Caption {j} (distance {best_distance}): Рис. {figure_number}")
                
                # Remove this caption from available captions to avoid double-linking
                del captions[figure_number]
            else:
                logger.warning(f"No caption found near image at position {i}")
        
        logger.info(f"Caption linking complete. Found {len(captions)} total captions, {len(image_fragments)} images in document order")
    
    def _link_citations_with_references(self, fragments: List) -> None:
        """Link citations in text [X] with bibliography at the end."""
        import re
        
        # Pattern for citations: [1], [2], etc.
        citation_pattern = re.compile(r'\[(\d+)\]')
        
        # First, find bibliography section (usually in last 20% of document)
        total_fragments = len(fragments)
        bibliography_start_idx = max(0, int(total_fragments * 0.8))  # Start looking from 80% into document
        
        bibliography = {}
        
        # Look for bibliography entries only in the last part of the document
        for i in range(bibliography_start_idx, total_fragments):
            fragment = fragments[i]
            if not hasattr(fragment, 'value') or not isinstance(fragment.value, str):
                continue
                
            text = fragment.value.strip()
            
            # Check if this looks like a bibliography entry (starts with number.)
            # Also check that it contains typical bibliography indicators
            bib_match = re.match(r'^(\d+)\.\s*(.+)', text)
            if bib_match and self._looks_like_bibliography_entry(bib_match.group(2)):
                ref_number = bib_match.group(1)
                ref_text = bib_match.group(2)
                
                bibliography[ref_number] = {
                    'text': ref_text,
                    'fragment_index': i,
                    'fragment_id': getattr(fragment, 'id', f'fragment_{i:03d}')
                }
                
                # Mark as bibliography entry
                if not hasattr(fragment, 'metadata'):
                    fragment.metadata = type('obj', (object,), {})()
                fragment.metadata.is_bibliography_entry = True
                fragment.metadata.reference_number = ref_number
        
        # Now find citations in text and link them
        for i, fragment in enumerate(fragments):
            if not hasattr(fragment, 'value') or not isinstance(fragment.value, str):
                continue
                
            text = fragment.value
            citations = citation_pattern.findall(text)
            
            if citations:
                if not hasattr(fragment, 'metadata'):
                    fragment.metadata = type('obj', (object,), {})()
                fragment.metadata.citations = []
                
                for citation_num in citations:
                    if citation_num in bibliography:
                        citation_info = {
                            'number': citation_num,
                            'reference_text': bibliography[citation_num]['text'],
                            'reference_fragment_id': bibliography[citation_num]['fragment_id']
                        }
                        fragment.metadata.citations.append(citation_info)
                        
                        # Add back-reference to bibliography entry
                        bib_fragment_idx = bibliography[citation_num]['fragment_index']
                        bib_fragment = fragments[bib_fragment_idx]
                        if not hasattr(bib_fragment, 'metadata'):
                            bib_fragment.metadata = type('obj', (object,), {})()
                        
                        if not hasattr(bib_fragment.metadata, 'cited_in'):
                            bib_fragment.metadata.cited_in = []
                        
                        bib_fragment.metadata.cited_in.append({
                            'fragment_id': getattr(fragment, 'id', f'fragment_{i:03d}'),
                            'fragment_text': text[:50] + '...' if len(text) > 50 else text
                        })
    
    def _looks_like_bibliography_entry(self, text: str) -> bool:
        """Check if text looks like a bibliography entry."""
        bibliography_indicators = [
            'http://', 'https://', 'www.', '.com', '.org', '.ru', '.net',
            'URL:', 'doi:', 'ISBN:', '//Журнал', '//М.:', '//СПб.:', 
            'Издательство', '– С.', 'С.', 'Т.', 'Vol.', 'No.', '№',
            'дата обращения', 'access date', 'retrieved', 'accessed'
        ]
        
        # Check if text contains any bibliography indicators
        text_lower = text.lower()
        for indicator in bibliography_indicators:
            if indicator.lower() in text_lower:
                return True
        
        # Also check length - bibliography entries are usually longer than table of contents entries
        return len(text) > 50
    
    def _detect_and_process_toc(self, fragments: List) -> None:
        """Detect and mark table of contents entries."""
        import re
        
        # Look for TOC patterns in early fragments (first 50% of document)
        toc_end_idx = min(len(fragments), int(len(fragments) * 0.5))
        
        # More comprehensive patterns that indicate TOC entries
        toc_patterns = [
            re.compile(r'^([А-Я][а-яё\s]+)\s*\.{2,}\s*(\d+)$'),  # "Введение....3"
            re.compile(r'^(\d+\.?\s*[А-Я][а-яё\s]+)\s*\.{2,}\s*(\d+)$'),  # "1. Анализ решений....4"
            re.compile(r'^(\d+\.\d+\.?\s*[А-Я][а-яё\s]+)\s*\.{2,}\s*(\d+)$'),  # "1.1. Подраздел....5"
            re.compile(r'^([А-Я][а-яё\s]+)\s+(\d+)$'),  # "Введение 3"
            re.compile(r'^(\d+\.?\s*[А-Я][а-яё\s]+)\s+(\d+)$'),  # "1. Анализ решений 4"
            re.compile(r'^([А-Я][а-яё\s,\.]+)\s*\t+(\d+)$'),  # "Введение\t\t3" (tab-separated)
            re.compile(r'^(\d+\.?\s*[А-Я][а-яё\s,\.]+)\s*\t+(\d+)$'),  # "1. Анализ\t\t4"
        ]
        
        toc_entries = []
        found_toc_section = False
        
        # First, look for TOC title to identify the section
        for i in range(min(30, len(fragments))):  # Check first 30 fragments for TOC title
            fragment = fragments[i]
            if hasattr(fragment, 'value') and isinstance(fragment.value, str):
                text = fragment.value.strip().lower()
                if text in ['содержание', 'оглавление', 'table of contents', 'contents']:
                    found_toc_section = True
                    # Mark this as TOC title
                    if not hasattr(fragment, 'metadata'):
                        fragment.metadata = type('obj', (object,), {})()
                    fragment.metadata.is_toc_title = True
                    break
        
        # Now look for TOC entries
        for i in range(min(toc_end_idx, len(fragments))):
            fragment = fragments[i]
            if not hasattr(fragment, 'value') or not isinstance(fragment.value, str):
                continue
                
            text = fragment.value.strip()
            
            # Skip if empty or too short
            if len(text) < 3:
                continue
            
            # Check if text matches TOC patterns
            for pattern in toc_patterns:
                match = pattern.match(text)
                if match:
                    title = match.group(1).strip()
                    page_num = match.group(2)
                    
                    # Additional validation - make sure it looks like TOC entry
                    if self._looks_like_toc_entry(title, page_num):
                        # Mark as TOC entry
                        if not hasattr(fragment, 'metadata'):
                            fragment.metadata = type('obj', (object,), {})()
                        
                        fragment.metadata.is_toc_entry = True
                        fragment.metadata.toc_title = title
                        fragment.metadata.toc_page = int(page_num)
                        
                        # Determine level based on numbering
                        if re.match(r'^\d+\.\d+', title):
                            fragment.metadata.toc_level = 2
                        elif re.match(r'^\d+\.', title):
                            fragment.metadata.toc_level = 1
                        else:
                            fragment.metadata.toc_level = 0  # Main sections like "Введение"
                        
                        toc_entries.append({
                            'title': title,
                            'page': int(page_num),
                            'level': fragment.metadata.toc_level,
                            'fragment_id': getattr(fragment, 'id', f'fragment_{i:03d}')
                        })
                        
                        logger.info(f"Found TOC entry: {title} -> page {page_num}")
                        break
        
        # If we found TOC entries, update page estimates for early fragments
        if toc_entries:
            logger.info(f"Found {len(toc_entries)} TOC entries")
            self._update_page_estimates_with_toc(fragments, toc_entries)
        else:
            logger.warning("No TOC entries found")
    
    def _looks_like_toc_entry(self, title: str, page_num: str) -> bool:
        """Check if this looks like a legitimate TOC entry."""
        try:
            page = int(page_num)
            # Page numbers should be reasonable (1-500)
            if page < 1 or page > 500:
                return False
            
            # Title should have some content
            if len(title.strip()) < 3:
                return False
            
            # Title should start with capital letter or number
            if not (title[0].isupper() or title[0].isdigit()):
                return False
            
            return True
        except ValueError:
            return False
    
    def _update_page_estimates_with_toc(self, fragments: List, toc_entries: List[Dict]) -> None:
        """Update page estimates based on table of contents information."""
        # Create a mapping of content titles to expected pages
        toc_mapping = {}
        for entry in toc_entries:
            toc_mapping[entry['title'].lower().strip()] = entry['page']
        
        # Go through fragments and update pages based on TOC mapping
        for i, fragment in enumerate(fragments):
            if not hasattr(fragment, 'value') or not isinstance(fragment.value, str):
                continue
            
            fragment_text = fragment.value.lower().strip()
            
            # Check if this fragment matches a TOC entry
            for toc_title, expected_page in toc_mapping.items():
                if toc_title in fragment_text or fragment_text in toc_title:
                    # Update the page for this fragment and nearby fragments
                    if hasattr(fragment, 'page'):
                        fragment.page = expected_page
                        
                        # Also update page metadata if it exists
                        if hasattr(fragment, 'metadata'):
                            fragment.metadata.expected_page_from_toc = expected_page
    
    def _is_header_style(self, style_name: str) -> bool:
        """Check if style indicates a header."""
        header_styles = [
            'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6',
            'Title', 'Subtitle'
        ]
        return style_name in header_styles
    
    def _get_header_level(self, style_name: str) -> int:
        """Get header level from style name."""
        if 'Title' in style_name:
            return 1
        elif 'Heading' in style_name:
            try:
                return int(style_name.split()[-1])
            except (ValueError, IndexError):
                return 2
        return 2
    
    def _is_list_item(self, paragraph) -> bool:
        """Check if paragraph is a list item."""
        # Simple check - can be enhanced
        return paragraph.style.name.startswith('List') if paragraph.style else False
    
    def _get_list_info(self, paragraph) -> Dict[str, Any]:
        """Get list information."""
        # Simplified implementation
        return {
            'type': 'unordered',
            'level': 0,
            'marker': '•'
        }
    
    def _has_rich_formatting(self, paragraph) -> bool:
        """Check if paragraph has rich formatting."""
        for run in paragraph.runs:
            if run.bold or run.italic or run.underline:
                return True
        return False
    
    def _extract_formatting(self, paragraph) -> Dict[str, Any]:
        """Extract detailed formatting information."""
        formatting = {
            'is_bold': False,
            'is_italic': False,
            'is_underline': False,
            'font_family': None,
            'font_size': None,
            'alignment': None,
            'text_color': None,
            'highlight_color': None,
            'is_subscript': False,
            'is_superscript': False,
            'is_strikethrough': False,
            'is_small_caps': False,
            'is_all_caps': False,
            'has_mixed_formatting': False,
            'bold_ratio': 0.0,
            'italic_ratio': 0.0,
            'left_indent': None,
            'right_indent': None,
            'first_line_indent': None,
            'space_before': None,
            'space_after': None,
            'line_spacing': None,
            'style_name': None,
            'style_type': None
        }
        
        # Collect formatting from all runs
        run_formats = []
        bold_count = 0
        italic_count = 0
        
        for run in paragraph.runs:
            run_format = {}
            
            # Basic formatting
            if run.bold:
                formatting['is_bold'] = True
                bold_count += 1
                run_format['bold'] = True
            if run.italic:
                formatting['is_italic'] = True
                italic_count += 1
                run_format['italic'] = True
            if run.underline:
                formatting['is_underline'] = True
                run_format['underline'] = True
            
            # Font properties
            if run.font.name:
                formatting['font_family'] = run.font.name
                run_format['font_name'] = run.font.name
            if run.font.size:
                formatting['font_size'] = run.font.size.pt
                run_format['font_size'] = run.font.size.pt
                
            # Color properties
            if run.font.color.rgb:
                formatting['text_color'] = str(run.font.color.rgb)
                run_format['color'] = str(run.font.color.rgb)
                
            # Additional properties (with safety checks)
            try:
                if hasattr(run.font, 'subscript') and run.font.subscript:
                    formatting['is_subscript'] = True
                if hasattr(run.font, 'superscript') and run.font.superscript:
                    formatting['is_superscript'] = True
                if hasattr(run.font, 'strike') and run.font.strike:
                    formatting['is_strikethrough'] = True
                if hasattr(run.font, 'small_caps') and run.font.small_caps:
                    formatting['is_small_caps'] = True
                if hasattr(run.font, 'all_caps') and run.font.all_caps:
                    formatting['is_all_caps'] = True
            except Exception:
                pass  # Some properties might not be available
                
            run_formats.append(run_format)
        
        # Calculate formatting ratios and mixed formatting
        total_runs = len(paragraph.runs)
        if total_runs > 0:
            formatting['bold_ratio'] = bold_count / total_runs
            formatting['italic_ratio'] = italic_count / total_runs
            formatting['has_mixed_formatting'] = len(set(str(r) for r in run_formats)) > 1
        
        # Paragraph-level formatting
        if paragraph.alignment:
            alignment_map = {
                WD_ALIGN_PARAGRAPH.LEFT: 'left',
                WD_ALIGN_PARAGRAPH.CENTER: 'center', 
                WD_ALIGN_PARAGRAPH.RIGHT: 'right',
                WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify'
            }
            formatting['alignment'] = alignment_map.get(paragraph.alignment, 'left')
        
        # Spacing and indentation (with safety checks)
        try:
            if paragraph.paragraph_format.left_indent:
                formatting['left_indent'] = paragraph.paragraph_format.left_indent.pt
            if paragraph.paragraph_format.right_indent:
                formatting['right_indent'] = paragraph.paragraph_format.right_indent.pt
            if paragraph.paragraph_format.first_line_indent:
                formatting['first_line_indent'] = paragraph.paragraph_format.first_line_indent.pt
            if paragraph.paragraph_format.space_before:
                formatting['space_before'] = paragraph.paragraph_format.space_before.pt
            if paragraph.paragraph_format.space_after:
                formatting['space_after'] = paragraph.paragraph_format.space_after.pt
            if paragraph.paragraph_format.line_spacing:
                formatting['line_spacing'] = paragraph.paragraph_format.line_spacing
        except Exception:
            pass  # Some properties might not be available
        
        # Style information
        if paragraph.style:
            formatting['style_name'] = paragraph.style.name
            try:
                formatting['style_type'] = str(paragraph.style.type)
            except Exception:
                pass
        
        return formatting
    
    def _process_drawing(self, drawing_element, docx_doc: "PythonDocxDocument", page: int = 1) -> Iterator:
        """Process drawing elements (images) in the document."""
        try:
            # Look for image elements in the drawing
            # Use simple xpath without namespaces parameter
            pics = drawing_element.xpath('.//pic:pic')
            for pic in pics:
                # Get image reference
                blips = pic.xpath('.//a:blip')
                if not blips:
                    continue
                
                # Get relationship ID
                embed_id = blips[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if not embed_id:
                    continue
                
                # Get image from relationships
                try:
                    image_part = docx_doc.part.related_parts[embed_id]
                    if hasattr(image_part, 'blob'):
                        # Create image fragment
                        from ...data.structures.fragment.image import ImageFragment
                        
                        # Generate unique ID for the image
                        image_id = f"image_{hash(embed_id) % 10000}"
                        
                        # Get image dimensions if available
                        width = None
                        height = None
                        try:
                            extents = pic.xpath('.//a:ext')
                            if extents:
                                width = int(extents[0].get('cx', 0)) // 9525  # Convert EMU to pixels
                                height = int(extents[0].get('cy', 0)) // 9525
                        except Exception:
                            pass
                        
                        # Create image fragment
                        try:
                            from PIL import Image
                            import io
                            
                            # Convert blob to PIL Image
                            image_data = io.BytesIO(image_part.blob)
                            pil_image = Image.open(image_data)
                            
                            image_fragment = ImageFragment(
                                value=pil_image,
                                format=image_part.content_type.split('/')[-1] if hasattr(image_part, 'content_type') else 'PNG',
                                page=page
                            )
                        except Exception as img_error:
                            logger.warning(f"Failed to convert image to PIL: {img_error}")
                            continue
                        
                        yield image_fragment
                        logger.info(f"Extracted image: {image_id} ({image_part.content_type if hasattr(image_part, 'content_type') else 'unknown'})")
                        
                except Exception as e:
                    logger.warning(f"Failed to extract image: {e}")
                    continue
                    
        except Exception as e:
            logger.warning(f"Error processing drawing element: {e}")
    
    def _extract_paragraph_images(self, paragraph, para_index: int, page: int = 1) -> Iterator:
        """Extract images from paragraph runs."""
        try:
            for run in paragraph.runs:
                # Check if run contains drawing elements
                # Use simple xpath without namespaces parameter
                drawings = run._element.xpath('.//w:drawing')
                for drawing in drawings:
                    yield from self._process_drawing(drawing, paragraph._parent, page)
        except Exception as e:
            logger.warning(f"Error extracting images from paragraph: {e}")
    
    def _process_paragraph_in_document_order(self, paragraph, para_index: int, page: int, 
                                           image_counter: int, processed_hashes: set, relationships: dict, file_path) -> Iterator:
        """
        Process paragraph content preserving exact document order.
        
        Handles all image types (DrawingML, VML, Objects, Embeds) and extracts them
        inline with text processing. Prevents over-fragmentation by processing
        complete paragraphs as single units while maintaining rich metadata.
        
        Args:
            paragraph: Python-docx Paragraph object
            para_index: Sequential paragraph index
            page: Estimated page number
            image_counter: Current image counter for unique naming
            processed_hashes: Set of processed image hashes for deduplication
            relationships: Document relationships mapping for image lookup
            file_path: Path to DOCX file for ZIP-based image extraction
            
        Yields:
            Tuple of (type, fragment, counter) for images or Fragment objects for text
        """
        try:
            # Check for ALL types of image elements in paragraph
            para_drawings = paragraph._element.xpath('.//w:drawing')
            para_picts = paragraph._element.xpath('.//w:pict')  # VML images
            
            # Also search for object elements that might contain images
            para_objects = paragraph._element.xpath('.//w:object')
            
            # Search for any element with relationship IDs that might point to images
            embed_elements = paragraph._element.xpath('.//*[@r:embed]')
            
            total_images = len(para_drawings) + len(para_picts) + len(para_objects) + len(embed_elements)
            if total_images > 0:
                logger.info(f"Paragraph {para_index}: found {total_images} images")
            
            # Get full paragraph text to avoid over-fragmentation
            full_para_text = paragraph.text.strip()
            
            # Process images if any (DrawingML, VML, Objects, Embeds)
            if para_drawings or para_picts or para_objects or embed_elements:
                # Process modern DrawingML images first
                for drawing_idx, drawing in enumerate(para_drawings):
                    # Extract rId from the drawing
                    blips = drawing.xpath('.//a:blip')
                    
                    if blips:
                        embed_attr = blips[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        
                        if embed_attr and embed_attr in relationships:
                            rel = relationships[embed_attr]
                            
                            # Extract image directly from DOCX ZIP archive using target_ref
                            if hasattr(rel, 'target_ref'):
                                try:
                                    from PIL import Image
                                    import io
                                    import zipfile
                                    from ...data.structures.fragment.image import ImageFragment
                                    
                                    # Open DOCX as ZIP and get image data
                                    with zipfile.ZipFile(str(file_path), 'r') as docx_zip:
                                        # target_ref is like "media/image3.png", we need "word/media/image3.png"
                                        image_path_in_zip = f"word/{rel.target_ref}"
                                        
                                        if image_path_in_zip in docx_zip.namelist():
                                            image_data = docx_zip.read(image_path_in_zip)
                                            pil_image = Image.open(io.BytesIO(image_data))
                                            
                                            # Check for duplicates
                                            img_hash = hash(pil_image.tobytes())
                                            if img_hash not in processed_hashes:
                                                processed_hashes.add(img_hash)
                                                image_counter += 1
                                                
                                                # Get format from target_ref
                                                format_ext = 'png'
                                                if '.' in rel.target_ref:
                                                    format_ext = rel.target_ref.split('.')[-1]
                                                
                                                image_fragment = ImageFragment(
                                                    value=pil_image,
                                                    format=format_ext,
                                                    page=page
                                                )
                                                
                                                # Add image metadata
                                                if not hasattr(image_fragment, 'metadata'):
                                                    image_fragment.metadata = type('obj', (object,), {})()
                                                
                                                image_fragment.metadata.image_type = "DrawingML"
                                                image_fragment.metadata.source_path = image_path_in_zip
                                                image_fragment.metadata.relationship_id = embed_attr
                                                image_fragment.metadata.image_index = image_counter
                                                image_fragment.metadata.width = pil_image.width
                                                image_fragment.metadata.height = pil_image.height
                                                image_fragment.metadata.mode = pil_image.mode
                                                image_fragment.metadata.file_size = len(image_data)
                                                image_fragment.metadata.paragraph_index = para_index
                                                
                                                logger.info(f"Extracted DrawingML image {image_counter}")
                                                yield ('image', image_fragment, image_counter)
                                        else:
                                            logger.warning(f"Image path {image_path_in_zip} not found in ZIP, available paths: {list(filter(lambda x: 'media' in x, docx_zip.namelist()))}")
                                        
                                except Exception as e:
                                    logger.warning(f"Failed to extract image from ZIP {embed_attr}: {e}")
                            else:
                                logger.warning(f"Relationship {embed_attr} has no target_ref")
                        else:
                            logger.warning(f"Embed attribute {embed_attr} not found in relationships")
                
                # Process legacy VML images  
                for pict_idx, pict in enumerate(para_picts):
                    # VML images use different structure - search without namespace prefix
                    imagedata_elements = pict.xpath('.//*[local-name()="imagedata"]')
                    
                    for imagedata in imagedata_elements:
                        # VML uses r:id instead of r:embed
                        rid = imagedata.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        if rid and rid in relationships:
                            rel = relationships[rid]
                            if hasattr(rel, 'target_ref'):
                                try:
                                    from PIL import Image
                                    import io
                                    import zipfile
                                    from ...data.structures.fragment.image import ImageFragment
                                    
                                    with zipfile.ZipFile(str(file_path), 'r') as docx_zip:
                                        image_path_in_zip = f"word/{rel.target_ref}"
                                        
                                        if image_path_in_zip in docx_zip.namelist():
                                            image_data = docx_zip.read(image_path_in_zip)
                                            pil_image = Image.open(io.BytesIO(image_data))
                                            
                                            img_hash = hash(pil_image.tobytes())
                                            if img_hash not in processed_hashes:
                                                processed_hashes.add(img_hash)
                                                image_counter += 1
                                                
                                                format_ext = 'png'
                                                if '.' in rel.target_ref:
                                                    format_ext = rel.target_ref.split('.')[-1]
                                                
                                                image_fragment = ImageFragment(
                                                    value=pil_image,
                                                    format=format_ext,
                                                    page=page
                                                )
                                                
                                                # Add VML image metadata
                                                if not hasattr(image_fragment, 'metadata'):
                                                    image_fragment.metadata = type('obj', (object,), {})()
                                                
                                                image_fragment.metadata.image_type = "VML"
                                                image_fragment.metadata.source_path = image_path_in_zip
                                                image_fragment.metadata.relationship_id = rid
                                                image_fragment.metadata.image_index = image_counter
                                                image_fragment.metadata.width = pil_image.width
                                                image_fragment.metadata.height = pil_image.height
                                                image_fragment.metadata.mode = pil_image.mode
                                                image_fragment.metadata.file_size = len(image_data)
                                                image_fragment.metadata.paragraph_index = para_index
                                                
                                                logger.info(f"Extracted VML image {image_counter}")
                                                yield ('image', image_fragment, image_counter)
                                                
                                        else:
                                            logger.warning(f"VML image path {image_path_in_zip} not found in ZIP")
                                except Exception as e:
                                    logger.warning(f"Failed to extract VML image from {rid}: {e}")
                        else:
                            logger.warning(f"VML r:id {rid} not found in relationships")
                
                # Process OLE objects that might contain images
                for obj_idx, obj in enumerate(para_objects):
                    # Objects might use different relationship patterns
                    obj_embeds = obj.xpath('.//*[@r:id]')
                    
                    for embed in obj_embeds:
                        rid = embed.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        if rid and rid in relationships:
                            rel = relationships[rid]
                            if hasattr(rel, 'target_ref') and 'media/' in rel.target_ref:
                                # Extract using same method as DrawingML
                                try:
                                    from PIL import Image
                                    import io
                                    import zipfile
                                    from ...data.structures.fragment.image import ImageFragment
                                    
                                    with zipfile.ZipFile(str(file_path), 'r') as docx_zip:
                                        image_path_in_zip = f"word/{rel.target_ref}"
                                        
                                        if image_path_in_zip in docx_zip.namelist():
                                            image_data = docx_zip.read(image_path_in_zip)
                                            pil_image = Image.open(io.BytesIO(image_data))
                                            
                                            img_hash = hash(pil_image.tobytes())
                                            if img_hash not in processed_hashes:
                                                processed_hashes.add(img_hash)
                                                image_counter += 1
                                                
                                                format_ext = 'png'
                                                if '.' in rel.target_ref:
                                                    format_ext = rel.target_ref.split('.')[-1]
                                                
                                                image_fragment = ImageFragment(
                                                    value=pil_image,
                                                    format=format_ext,
                                                    page=page
                                                )
                                                
                                                # Add object image metadata
                                                if not hasattr(image_fragment, 'metadata'):
                                                    image_fragment.metadata = type('obj', (object,), {})()
                                                
                                                image_fragment.metadata.image_type = "Object"
                                                image_fragment.metadata.source_path = image_path_in_zip
                                                image_fragment.metadata.relationship_id = rid
                                                image_fragment.metadata.image_index = image_counter
                                                image_fragment.metadata.width = pil_image.width
                                                image_fragment.metadata.height = pil_image.height
                                                image_fragment.metadata.mode = pil_image.mode
                                                image_fragment.metadata.file_size = len(image_data)
                                                image_fragment.metadata.paragraph_index = para_index
                                                
                                                logger.info(f"Extracted object image {image_counter}")
                                                yield ('image', image_fragment, image_counter)
                                                
                                        else:
                                            logger.warning(f"Object image path {image_path_in_zip} not found in ZIP")
                                except Exception as e:
                                    logger.warning(f"Failed to extract object image from {rid}: {e}")
                
                # Process any other embedded elements with r:embed attributes
                for embed_idx, embed in enumerate(embed_elements):
                    if embed in para_drawings or embed in para_picts or embed in para_objects:
                        continue  # Already processed above
                    
                    rid = embed.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed') or \
                          embed.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    
                    if rid and rid in relationships:
                        rel = relationships[rid]
                        if hasattr(rel, 'target_ref') and 'media/' in rel.target_ref:
                            # Extract using same method
                            try:
                                from PIL import Image
                                import io
                                import zipfile
                                from ...data.structures.fragment.image import ImageFragment
                                
                                with zipfile.ZipFile(str(file_path), 'r') as docx_zip:
                                    image_path_in_zip = f"word/{rel.target_ref}"
                                    
                                    if image_path_in_zip in docx_zip.namelist():
                                        image_data = docx_zip.read(image_path_in_zip)
                                        pil_image = Image.open(io.BytesIO(image_data))
                                        
                                        img_hash = hash(pil_image.tobytes())
                                        if img_hash not in processed_hashes:
                                            processed_hashes.add(img_hash)
                                            image_counter += 1
                                            
                                            format_ext = 'png'
                                            if '.' in rel.target_ref:
                                                format_ext = rel.target_ref.split('.')[-1]
                                            
                                            image_fragment = ImageFragment(
                                                value=pil_image,
                                                format=format_ext,
                                                page=page
                                            )
                                            
                                            # Add generic embed image metadata
                                            if not hasattr(image_fragment, 'metadata'):
                                                image_fragment.metadata = type('obj', (object,), {})()
                                            
                                            image_fragment.metadata.image_type = "Embed"
                                            image_fragment.metadata.source_path = image_path_in_zip
                                            image_fragment.metadata.relationship_id = rid
                                            image_fragment.metadata.image_index = image_counter
                                            image_fragment.metadata.width = pil_image.width
                                            image_fragment.metadata.height = pil_image.height
                                            image_fragment.metadata.mode = pil_image.mode
                                            image_fragment.metadata.file_size = len(image_data)
                                            image_fragment.metadata.paragraph_index = para_index
                                            
                                            logger.info(f"Extracted embed image {image_counter}")
                                            yield ('image', image_fragment, image_counter)
                                            
                                    else:
                                        logger.warning(f"Generic embed image path {image_path_in_zip} not found in ZIP")
                            except Exception as e:
                                logger.warning(f"Failed to extract generic embed image from {rid}: {e}")
            
            # Process text as single paragraph, but check for header styles
            if full_para_text:
                # Extract detailed formatting first
                detailed_formatting = self._extract_formatting(paragraph)
                
                # Determine paragraph type based on style
                style_name = paragraph.style.name if paragraph.style else ""
                
                # Check if it's a header
                if self._is_header_style(style_name):
                    level = self._get_header_level(style_name)
                    if level == 1:
                        fragment = TitleFragment(
                            value=full_para_text,
                            id=f"title_{para_index}",
                            level=level,
                            page=page
                        )
                    else:
                        fragment = HeaderFragment(
                            value=full_para_text,
                            id=f"header_{para_index}_{level}",
                            level=level,
                            page=page
                        )
                
                # Check if it's a list item
                elif self._is_list_item(paragraph):
                    list_info = self._get_list_info(paragraph)
                    fragment = ListItemFragment(
                        value=full_para_text,
                        id=f"list_item_{para_index}",
                        list_level=list_info['level'],
                        list_type=list_info['type'],
                        marker=list_info['marker'],
                        page=page
                    )
                
                # Regular paragraph
                else:
                    fragment = ParagraphFragment(
                        value=full_para_text,
                        id=f"para_{para_index}",
                        page=page
                    )
                
                # Add comprehensive metadata to ALL fragments
                if not hasattr(fragment, 'metadata'):
                    fragment.metadata = type('obj', (object,), {})()
                
                # Add detailed formatting metadata
                for key, value in detailed_formatting.items():
                    if value is not None:
                        setattr(fragment.metadata, key, value)
                
                # Add paragraph-level metadata
                fragment.metadata.style_name = style_name
                fragment.metadata.paragraph_index = para_index
                fragment.metadata.text_length = len(full_para_text)
                fragment.metadata.word_count = len(full_para_text.split())
                fragment.metadata.has_images = total_images > 0
                fragment.metadata.run_count = len(paragraph.runs)
                
                # Add document structure metadata if available
                if hasattr(paragraph, '_element'):
                    try:
                        # Extract alignment from paragraph properties
                        pPr = paragraph._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                        if pPr is not None:
                            jc = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc')
                            if jc is not None:
                                fragment.metadata.paragraph_alignment = jc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    except Exception:
                        pass
                
                yield fragment
                        
        except Exception as e:
            logger.warning(f"Error processing paragraph in document order: {e}")
            # Fallback to normal paragraph processing
            yield from self._process_paragraph(paragraph, para_index, page)
    
    def _process_table_with_images(self, table, page: int, image_counter: int, processed_hashes: set, relationships: dict, file_path) -> Iterator:
        """Process table and extract any images found within."""
        try:
            logger.info(f"Processing table with {len(table.rows)} rows for images")
            
            # Check each cell for images
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Process each paragraph in the cell
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        for result in self._process_paragraph_in_document_order(
                            paragraph, f"table_r{row_idx}_c{cell_idx}_p{para_idx}", page, 
                            image_counter, processed_hashes, relationships, file_path
                        ):
                            if isinstance(result, tuple) and result[0] == 'image':
                                image_counter = result[2]  # Update counter
                                yield result[1]  # Yield image fragment
                            else:
                                yield result  # Yield text fragment
                                
        except Exception as e:
            logger.warning(f"Error processing table with images: {e}")
            # Fallback to normal table processing
            yield from self._process_table(table, page)
    
    def _extract_image_from_drawing(self, drawing_element, relationships: dict, 
                                  image_counter: int, processed_hashes: set, page: int):
        """Extract single image from drawing element using relationships."""
        try:
            # Look for image references in drawing
            blips = drawing_element.xpath('.//a:blip')
            logger.debug(f"Found {len(blips)} blip elements in drawing")
            if not blips:
                return None
            
            # Get the relationship ID
            embed_attr = blips[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            logger.debug(f"Embed attribute: {embed_attr}, Available relationships: {list(relationships.keys())}")
            if not embed_attr or embed_attr not in relationships:
                logger.debug(f"Embed attribute {embed_attr} not found in relationships")
                return None
            
            # Get image from relationships
            rel = relationships[embed_attr]
            if not hasattr(rel, 'blob'):
                return None
            
            # Create PIL image
            from PIL import Image
            import io
            
            image_data = io.BytesIO(rel.blob)
            pil_image = Image.open(image_data)
            
            # Check for duplicates
            img_hash = hash(pil_image.tobytes())
            if img_hash in processed_hashes:
                return None
            
            processed_hashes.add(img_hash)
            image_counter += 1
            
            # Create image fragment
            from ...data.structures.fragment.image import ImageFragment
            image_fragment = ImageFragment(
                value=pil_image,
                format=rel.content_type.split('/')[-1] if hasattr(rel, 'content_type') else 'png',
                page=page
            )
            
            logger.info(f"Extracted image in document order: image_{image_counter}")
            return image_fragment, image_counter
            
        except Exception as e:
            logger.warning(f"Error extracting image from drawing: {e}")
            return None
    
    def _process_drawing_with_counter(self, drawing_element, docx_doc, page: int, 
                                    image_counter: int, processed_hashes: set) -> Iterator:
        """Process drawing elements with counter for proper ordering."""
        try:
            pics = drawing_element.xpath('.//pic:pic')
            for pic in pics:
                blips = pic.xpath('.//a:blip')
                if not blips:
                    continue
                
                embed_id = blips[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if not embed_id:
                    continue
                
                try:
                    image_part = docx_doc.part.related_parts[embed_id]
                    if hasattr(image_part, 'blob'):
                        from ...data.structures.fragment.image import ImageFragment
                        from PIL import Image
                        import io
                        
                        # Convert blob to PIL Image
                        image_data = io.BytesIO(image_part.blob)
                        pil_image = Image.open(image_data)
                        
                        # Create hash for deduplication
                        img_hash = hash(pil_image.tobytes())
                        if img_hash not in processed_hashes:
                            processed_hashes.add(img_hash)
                            image_counter += 1
                            
                            image_fragment = ImageFragment(
                                value=pil_image,
                                format=image_part.content_type.split('/')[-1] if hasattr(image_part, 'content_type') else 'PNG',
                                page=page
                            )
                            
                            logger.info(f"Extracted image from paragraph in order: image_{image_counter}")
                            yield image_fragment, image_counter
                            
                except Exception as e:
                    logger.warning(f"Failed to extract image: {e}")
                    continue
                    
        except Exception as e:
            logger.warning(f"Error processing drawing element: {e}")
